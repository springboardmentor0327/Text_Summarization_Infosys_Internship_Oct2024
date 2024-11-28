from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords
import nltk
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from sumy.summarizers.text_rank import TextRankSummarizer
from sumy.summarizers.lex_rank import LexRankSummarizer
from sumy.summarizers.lsa import LsaSummarizer
import docx


nltk.download('stopwords')
nltk.download('punkt')
nltk.download('punkt_tab')

stopwords1 = set(stopwords.words("english"))
# Define functions for reading files
def read_pdf(file):
    text = ""
    with fitz.open(file.name) as pdf:
        for page in pdf:
            text += page.get_text()
    return text

def read_docx(file):
    text = ""
    doc = docx.Document(file.name)
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def process_file(file):
    if file.name.endswith('.pdf'):
        return read_pdf(file)
    elif file.name.endswith('.docx'):
        return read_docx(file)
    else:
        return "Unsupported file type."

def format_summary(sentences, summary_style):
    if summary_style == "Detailed":
        return ' '.join(sentences)
    elif summary_style == "Elaborated":
        return ' '.join(sentences)
    elif summary_style == "Bullet Points":
        return '\n'.join(f"- {sentence}" for sentence in sentences)

def freq_method(text, word_count, summary_style):
    words = word_tokenize(text)

    # Frequency Table Creation
    freqTable = {}
    for word in words:
        word = word.lower()
        if word in stopwords1:
            continue
        freqTable[word] = freqTable.get(word, 0) + 1

    sentences = sent_tokenize(text)
    sentenceValue = {}

    for sentence in sentences:
        for word, freq in freqTable.items():
            if word in sentence.lower():
                sentenceValue[sentence] = sentenceValue.get(sentence, 0) + freq

    # Select sentences until word count is met
    selected_sentences = []
    word_total = 0
    for sentence in sorted(sentenceValue, key=sentenceValue.get, reverse=True):
        selected_sentences.append(sentence)
        word_total += len(word_tokenize(sentence))
        if word_total >= word_count:  # Target within range
            break

    return format_summary(selected_sentences, summary_style)

def sumy_method(text, word_count, summary_style):
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = TextRankSummarizer()
    sentences = [str(sentence) for sentence in summarizer(parser.document, 10)]

    # Select sentences to meet word count target
    selected_sentences = []
    word_total = 0
    for sentence in sentences:
        selected_sentences.append(sentence)
        word_total += len(word_tokenize(sentence))
        if word_total >= word_count - 10:
            break

    return format_summary(selected_sentences, summary_style)

def lex_rank(text, word_count, summary_style):
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer = LexRankSummarizer()
    sentences = [str(sentence) for sentence in summarizer(parser.document, 10)]

    # Select sentences until word count target
    selected_sentences = []
    word_total = 0
    for sentence in sentences:
        selected_sentences.append(sentence)
        word_total += len(word_tokenize(sentence))
        if word_total >= word_count - 10:
            break

    return format_summary(selected_sentences, summary_style)

def lsa(text, word_count, summary_style):
    parser = PlaintextParser.from_string(text, Tokenizer("english"))
    summarizer_lsa = LsaSummarizer()
    sentences = [str(sentence) for sentence in summarizer_lsa(parser.document, 10)]

    # Select sentences until word count target
    selected_sentences = []
    word_total = 0
    for sentence in sentences:
        selected_sentences.append(sentence)
        word_total += len(word_tokenize(sentence))
        if word_total >= word_count - 10:
            break

    return format_summary(selected_sentences, summary_style)

# Abstractive Summarization Functions (BART, T5)
from transformers import BartForConditionalGeneration, BartTokenizer, T5ForConditionalGeneration, T5Tokenizer

def bart_summarization(text, word_count, summary_style):
    model = BartForConditionalGeneration.from_pretrained("facebook/bart-large-cnn")
    tokenizer = BartTokenizer.from_pretrained("facebook/bart-large-cnn")

    max_length = min(word_count + 10, 300)  # Cap to 300 words for coherence
    inputs = tokenizer([text], max_length=1024, return_tensors='pt', truncation=True)
    summary_ids = model.generate(inputs['input_ids'], max_length=max_length, length_penalty=2.0, num_beams=4, early_stopping=True)
    summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)

    # Format as bullet points if selected
    if summary_style == "Bullet Points":
        summary = '\n- ' + summary.replace(". ", ".\n- ")
    return summary

def t5_summarization(text, word_count, summary_style):
    model = T5ForConditionalGeneration.from_pretrained("t5-base")
    tokenizer = T5Tokenizer.from_pretrained("t5-base")

    max_length = min(word_count + 10, 300)
    input_text = "summarize: " + text
    inputs = tokenizer.encode(input_text, max_length=512, return_tensors='pt', truncation=True)
    summary_ids = model.generate(inputs, max_length=max_length, length_penalty=2.0, num_beams=4, early_stopping=True)
    summary = tokenizer.decode(summary_ids[0], skip_special_tokens=True)

    if summary_style == "Bullet Points":
        summary = '\n- ' + summary.replace(". ", ".\n- ")
    return summary

# LLM APPROACH

import os
userdata = {'GOOGLE_API_KEY': '<Enter your Gemini API key here'}
os.environ["GOOGLE_API_KEY"] = userdata.get('GOOGLE_API_KEY')
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate

def load_llm(model="gemini-1.5-flash"):
    if model == "gemini-1.5-pro":
        llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-pro",
            temperature=0,
            max_tokens=None,
            timeout=None,
            max_retries=2)
        return llm
    elif model == "gemini-1.5-flash":
        llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash",
            temperature=0,
            max_tokens=None,
            timeout=None,
            max_retries=2)
        return llm
    else:
        raise ValueError("Invalid model name")

def get_prompt_template():
    # Define prompt with dynamic word count input
    prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "Write a concise summary of the following in {num_words} words:\n\n",
            ),
            ("human", "{context}")
        ]
    )
    return prompt

def summarize_text(text, num_words=50, model="gemini-1.5-flash"):
    # Load LLM
    llm = load_llm(model)
    # Get Prompt Template
    prompt = get_prompt_template()
    # Instantiate chain
    chain = prompt | llm
    # Invoke chain with the desired word count
    result = chain.invoke({
        "context": text,
        "num_words": num_words
    })
    # Return result
    return result.content

####### Iterative Refinement ######
import operator
from typing import List, Literal, TypedDict

from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableConfig
from langgraph.constants import Send
from langgraph.graph import END, START, StateGraph
from nltk.tokenize import sent_tokenize
import nest_asyncio
import asyncio


# Initial summary
summarize_prompt = ChatPromptTemplate(
    [
        ("human", "Write a concise summary of the following: {context}"),
    ]
)
llm = load_llm()
initial_summary_chain = summarize_prompt | llm | StrOutputParser()

# Refining the summary with new docs
refine_template = """
Produce a final summary.

Existing summary up to this point:
{existing_answer}

New context:
------------
{context}
------------

Given the new context, refine the original summary.
"""
refine_prompt = ChatPromptTemplate([("human", refine_template)])

refine_summary_chain = refine_prompt | llm | StrOutputParser()


# We will define the state of the graph to hold the document
# contents and summary. We also include an index to keep track
# of our position in the sequence of documents.
class State(TypedDict):
    contents: List[str]
    index: int
    summary: str


# We define functions for each node, including a node that generates
# the initial summary:
async def generate_initial_summary(state: State, config: RunnableConfig):
    summary = await initial_summary_chain.ainvoke(
        state["contents"][0],
        config,
    )
    return {"summary": summary, "index": 1}


# And a node that refines the summary based on the next document
async def refine_summary(state: State, config: RunnableConfig):
    content = state["contents"][state["index"]]
    summary = await refine_summary_chain.ainvoke(
        {"existing_answer": state["summary"], "context": content},
        config,
    )

    return {"summary": summary, "index": state["index"] + 1}


# Here we implement logic to either exit the application or refine
# the summary.
def should_refine(state: State) -> Literal["refine_summary", END]:
    if state["index"] >= len(state["contents"]):
        return END
    else:
        return "refine_summary"


graph = StateGraph(State)
graph.add_node("generate_initial_summary", generate_initial_summary)
graph.add_node("refine_summary", refine_summary)

graph.add_edge(START, "generate_initial_summary")
graph.add_conditional_edges("generate_initial_summary", should_refine)
graph.add_conditional_edges("refine_summary", should_refine)
app = graph.compile()



async def summarize_documents(user_content: List[str]):
    summary = None
    async for step in app.astream(
        {"contents": user_content}, stream_mode="values"
    ):
        if step.get("summary"):
            summary = step["summary"]
    return summary


###### MAP REDUCE ######

from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate

map_prompt = ChatPromptTemplate.from_messages(
    [("human", "Write a concise summary of the following:\\n\\n{context}")]
)

map_chain = map_prompt | llm | StrOutputParser()

reduce_template = """
The following is a set of summaries:
{docs}
Take these and distill it into a final, consolidated summary
of the main themes.
"""

reduce_prompt = ChatPromptTemplate([("human", reduce_template)])

reduce_chain = reduce_prompt | llm | StrOutputParser()

import operator
from typing import Annotated, List, Literal, TypedDict

from langchain.chains.combine_documents.reduce import (
    acollapse_docs,
    split_list_of_docs,
)
from langchain_core.documents import Document
from langgraph.constants import Send
from langgraph.graph import END, START, StateGraph

token_max = 1000


def length_function(documents: List[Document]) -> int:
    """Get number of tokens for input contents."""
    return sum(llm.get_num_tokens(doc.page_content) for doc in documents)


# This will be the overall state of the main graph.
# It will contain the input document contents, corresponding
# summaries, and a final summary.
class OverallState(TypedDict):
    # Notice here we use the operator.add
    # This is because we want combine all the summaries we generate
    # from individual nodes back into one list - this is essentially
    # the "reduce" part
    contents: List[str]
    summaries: Annotated[list, operator.add]
    collapsed_summaries: List[Document]
    final_summary: str


# This will be the state of the node that we will "map" all
# documents to in order to generate summaries
class SummaryState(TypedDict):
    content: str


# Here we generate a summary, given a document
async def generate_summary(state: SummaryState):
    response = await map_chain.ainvoke(state["content"])
    return {"summaries": [response]}


# Here we define the logic to map out over the documents
# We will use this an edge in the graph
def map_summaries(state: OverallState):
    # We will return a list of `Send` objects
    # Each `Send` object consists of the name of a node in the graph
    # as well as the state to send to that node
    return [
        Send("generate_summary", {"content": content}) for content in state["contents"]
    ]


def collect_summaries(state: OverallState):
    return {
        "collapsed_summaries": [Document(summary) for summary in state["summaries"]]
    }


# Add node to collapse summaries
async def collapse_summaries(state: OverallState):
    doc_lists = split_list_of_docs(
        state["collapsed_summaries"], length_function, token_max
    )
    results = []
    for doc_list in doc_lists:
        results.append(await acollapse_docs(doc_list, reduce_chain.ainvoke))

    return {"collapsed_summaries": results}


# This represents a conditional edge in the graph that determines
# if we should collapse the summaries or not
def should_collapse(
    state: OverallState,
) -> Literal["collapse_summaries", "generate_final_summary"]:
    num_tokens = length_function(state["collapsed_summaries"])
    if num_tokens > token_max:
        return "collapse_summaries"
    else:
        return "generate_final_summary"


# Here we will generate the final summary
async def generate_final_summary(state: OverallState):
    response = await reduce_chain.ainvoke(state["collapsed_summaries"])
    return {"final_summary": response}


# Construct the graph
# Nodes:
graph = StateGraph(OverallState)
graph.add_node("generate_summary", generate_summary)  # same as before
graph.add_node("collect_summaries", collect_summaries)
graph.add_node("collapse_summaries", collapse_summaries)
graph.add_node("generate_final_summary", generate_final_summary)

# Edges:
graph.add_conditional_edges(START, map_summaries, ["generate_summary"])
graph.add_edge("generate_summary", "collect_summaries")
graph.add_conditional_edges("collect_summaries", should_collapse)
graph.add_conditional_edges("collapse_summaries", should_collapse)
graph.add_edge("generate_final_summary", END)

app = graph.compile()

# Gradio UI
css = """
h1 {
    margin-top: 2rem;
    font-size: 2rem;
    text-align: center;
}
"""
def extractive(text, method, length, summary_style):
    if method == "Frequency":
        return freq_method(text, length, summary_style)
    elif method == "Sumy":
        return sumy_method(text, length, summary_style)
    elif method == "LexRank":
        return lex_rank(text, length, summary_style)
    elif method == "LDA":
        return lsa(text, length, summary_style)
    else:
        return "Method not implemented."

def abstractive(text, method, length, summary_style):
    if method == "BART":
        return bart_summarization(text, length, summary_style)
    elif method == "T5":
        return t5_summarization(text, length, summary_style)
    elif method == "LLM":
        # Adjusting word count based on length choice
        #word_count = 50 if length == "Short" else 100 if length == "Medium" else 150
        return summarize_text(text, num_words=length,model="gemini-1.5-flash")
    else:
        return "Method not implemented."

def LLM(text, method, word_count, summary_style):
    if method == "Basic LLM":
        return summarize_text(text, num_words=word_count,model="gemini-1.5-flash")
    elif method == "Iterative Refinement":
        # Use nest_asyncio to allow running async code in non-async environments
        nest_asyncio.apply()
        
        # Convert text to sentences
        user_content = sent_tokenize(text)
        
        # Run the async function and get the result
        async def run_summary():
            return await summarize_documents(user_content)
        
        # Use asyncio to run the coroutine
        summary = asyncio.run(run_summary())
        return summary
    elif method == "Map Reduce":
      user_content = sent_tokenize(text)
        # Convert sentences into Document objects for processing
      split_docs = [Document(content) for content in user_content]
        
      async def map_reduce_summary():
          # Run the Map Reduce graph
          async for step in app.astream(
              {"contents": [doc.page_content for doc in split_docs]},
              {"recursion_limit": 10},
          ):
              if "final_summary" in step:
                  return step["final_summary"]

        # Execute the async function
      nest_asyncio.apply()  # Allow async in environments like Jupyter/Colab
      loop = asyncio.get_event_loop()
      final_summary = loop.run_until_complete(map_reduce_summary())
      return final_summary


import gradio as gr

input_text = gr.Textbox(label="Input Text", lines=10)
file_input = gr.File(label="Upload Document (PDF/Word)")

with gr.Blocks(title="Summarizer App", css=css) as demo:
    gr.Markdown("# Summarizer App")

    with gr.Tabs():
        with gr.TabItem("Extractive"):
            method_dropdown = gr.Dropdown(choices=["Frequency", "Sumy", "LexRank", "LDA"], label="Select Extractive Method")
            length_slider = gr.Slider(minimum=50, maximum=300, step=10, label="Set Summary Word Count")
            style_dropdown = gr.Dropdown(choices=["Detailed", "Bullet Points"], label="Select Summary Style")
            output = gr.Textbox(label="Summary", lines=10)

            def extractive_interface(text, method, word_count, style, file):
                if file is not None:
                    text = process_file(file)
                return extractive(text, method, word_count, style)

            gr.Interface(fn=extractive_interface,
                         inputs=[input_text, method_dropdown, length_slider, style_dropdown, file_input],
                         outputs=output)

        with gr.TabItem("Abstractive"):
            method_dropdown = gr.Dropdown(choices=["T5","BART","LLM"], label="Select Abstractive Method")
            length_slider = gr.Slider(minimum=50, maximum=300, step=10, label="Set Summary Word Count")
            style_dropdown = gr.Dropdown(choices=["Detailed", "Bullet Points"], label="Select Summary Style")
            output = gr.Textbox(label="Summary", lines=10)

            def abstractive_interface(text, method, word_count, style, file):
                if file is not None:
                    text = process_file(file)
                return abstractive(text, method, word_count, style)

            gr.Interface(fn=abstractive_interface,
                         inputs=[input_text, method_dropdown, length_slider, style_dropdown, file_input],
                         outputs=output)
            
        
        with gr.TabItem("LLM"):
            method_dropdown = gr.Dropdown(choices=["Basic LLM", "Iterative Refinement", "Map Reduce"], label="Select LLM Method")
            length_slider = gr.Slider(minimum=50, maximum=300, step=10, label="Set Summary Word Count")
            style_dropdown = gr.Dropdown(choices=["Detailed", "Bullet Points"], label="Select Summary Style")
            output = gr.Textbox(label="Summary", lines=10)

            def LLM_interface(text, method, word_count, style, file):
                if file is not None:
                    text = process_file(file)
                return LLM(text, method, word_count, style)

            gr.Interface(fn=LLM_interface,
                         inputs=[input_text, method_dropdown, length_slider, style_dropdown, file_input],
                         outputs=output)

# Launch the app
demo.launch()
